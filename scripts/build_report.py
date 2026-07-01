from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "说明文档.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_east_asia_font(run, font_name: str = "Microsoft YaHei") -> None:
    run.font.name = "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in (
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("面向计算机考研《数据结构》的 RAG 智能问答 Agent")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string("0B2545")
    set_east_asia_font(run)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("说明文档：系统功能、架构设计、知识库流程与运行方式")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("586575")
    set_east_asia_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        set_cell_shading(hdr[idx], "F2F4F7")
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(hdr[idx])
        for p in hdr[idx].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = text
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])


def build_doc() -> None:
    doc = Document()
    style_document(doc)
    add_title(doc)

    doc.add_heading("1. 系统概述", level=1)
    doc.add_paragraph(
        "本系统实现一个面向计算机考研《数据结构》科目的智能问答 Agent。系统基于本地知识库完成检索增强问答，"
        "同时提供数据结构代码生成能力。交互界面采用 Streamlit，适合课程考核演示、截图和本地运行。"
    )
    add_bullets(
        doc,
        [
            "支持概念解释、对比分析和代码生成三类核心任务。",
            "支持当前会话内的基础多轮对话，可结合上一轮问题理解追问。",
            "回答知识库问题时展示引用依据，不在资料不足时伪造来源。",
            "界面展示工具调用日志和检索命中文档、片段、相似度分数。",
        ],
    )

    doc.add_heading("2. 架构设计", level=1)
    doc.add_paragraph(
        "项目采用轻量分层架构。Streamlit 负责页面交互，Agent 负责问题路由，RAG 模块负责检索与问答，"
        "代码生成模块负责算法代码输出，导入模块负责资料读取、切分和向量库构建。"
    )
    add_table(
        doc,
        ["模块", "职责"],
        [
            ["app.py", "提供问题输入区、答案展示区、知识库上传区、工具日志区和检索结果区。"],
            ["src/agent.py", "根据问题类型自动选择 rag_qa 或 code_gen，并处理基础追问上下文。"],
            ["src/rag.py", "执行知识库检索、构造提示词、调用 OpenAI 兼容 API 并追加引用依据。"],
            ["src/codegen.py", "生成数据结构范围内的代码、核心思路和复杂度说明。"],
            ["src/ingest.py", "读取 txt、md、pdf 资料，清洗切分文本，生成 TF-IDF 向量并写入 FAISS 索引。"],
            ["src/config.py", "读取 .env 配置，统一管理模型、路径、检索阈值和分块参数。"],
        ],
    )

    doc.add_heading("3. 知识库构建流程", level=1)
    add_numbered(
        doc,
        [
            "将《数据结构》教材、PPT、笔记、真题解析或算法模板放入 materials 目录。",
            "系统读取 .txt、.md、.pdf 文件，其中 PDF 逐页提取文本。",
            "文本经过空白清洗、Markdown 标题分段和滑动窗口切分。",
            "使用中文友好的字符 n-gram TF-IDF 向量化，向量归一化后存入 FAISS 本地索引。",
            "用户提问时检索 Top-K 片段，展示来源、位置、片段内容和相似度分数。",
        ],
    )

    doc.add_heading("4. 工具调用逻辑", level=1)
    add_table(
        doc,
        ["工具", "触发条件", "输出"],
        [
            ["rag_qa", "概念解释、对比分析、复杂度解释等知识库类问题。", "基于检索片段生成回答，并附引用依据。"],
            ["code_gen", "出现代码、实现、程序、函数、C++、Python、Java 等代码生成意图。", "输出核心思路、完整代码、时间复杂度、空间复杂度和易错点。"],
        ],
    )
    doc.add_paragraph(
        "若检索结果相似度低于阈值，系统会提示当前知识库未找到可靠依据，不继续编造答案。"
    )

    doc.add_heading("5. 运行方式", level=1)
    add_numbered(
        doc,
        [
            "执行 pip install -r requirements.txt 安装依赖。",
            "复制 .env.example 为 .env，并填写 OPENAI_API_KEY、OPENAI_BASE_URL、OPENAI_MODEL。",
            "将知识库资料放入 materials 目录，或在页面侧边栏上传资料。",
            "执行 streamlit run app.py 启动系统。",
            "在页面点击“重建本地知识库”，然后输入问题进行测试。",
        ],
    )

    doc.add_heading("6. 测试样例", level=1)
    add_bullets(
        doc,
        [
            "概念解释：什么是栈？它有哪些基本操作？",
            "对比分析：顺序表和链表有什么区别？它们分别适合什么场景？",
            "代码生成：请用 C++ 实现单链表逆置，并说明时间复杂度和空间复杂度。",
            "多轮对话：先问“什么是快速排序？”，再问“它的时间复杂度是多少？”。",
            "低相关测试：询问知识库外问题，系统应提示缺少可靠依据。",
        ],
    )

    doc.add_heading("7. 效果截图位置", level=1)
    doc.add_paragraph("运行 Streamlit 后建议截取以下界面，粘贴到提交材料或答辩 PPT 中：")
    add_bullets(
        doc,
        [
            "主页面：问题输入区与答案展示区。",
            "工具日志：展示本轮调用 rag_qa 或 code_gen。",
            "检索结果：展示命中文档、片段内容、相似度分数和来源位置。",
            "知识库上传：展示上传资料并重建向量库的过程。",
        ],
    )

    doc.add_heading("8. 交付清单", level=1)
    add_bullets(
        doc,
        [
            "源代码目录：app.py、src、materials、vector_store、requirements.txt、.env.example。",
            "说明文档：说明文档.docx。",
            "最终压缩包：数据结构RAG智能问答Agent.zip。",
        ],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
